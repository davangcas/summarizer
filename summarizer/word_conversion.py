"""Extracción directa de contenido Word usando librerías Python."""

from __future__ import annotations

from pathlib import Path
from docx import Document


def extract_docx_text(docx_path: Path) -> str:
    """
    Extrae texto de un .docx incluyendo párrafos y tablas en formato simple.
    """
    doc = Document(str(docx_path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cols = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cols):
                rows.append(" | ".join(cols))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)
